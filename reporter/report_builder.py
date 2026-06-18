"""
Report Builder Module
使用 python-docx 生成专业中文地质勘探报告。
"""

from pathlib import Path
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .categories import SearchResult, get_all_categories
from .geocoder import LocationContext


class ReportBuilderError(Exception):
    pass


# 无公开数据的占位文本特征——命中则该条不展示（不输出"暂缺/来源"等占位句）
_NO_DATA_MARKERS = ("暂缺", "暂无", "无数据", "无相关", "未获取", "未检索", "n/a", "N/A", "待补充", "缺失")


def _is_no_data(value: str) -> bool:
    if not value or not str(value).strip():
        return True
    v = str(value)
    return any(m in v for m in _NO_DATA_MARKERS)


class ReportBuilder:
    """使用 python-docx 构建中文专业报告"""

    # 中文字体设置（GB/T 9704 标准）
    FONT_SONGTI = "宋体"
    FONT_HEITI = "黑体"
    FONT_KAITI = "楷体"

    # 颜色定义
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_GRAY = RGBColor(128, 128, 128)
    COLOR_LIGHT_GRAY = RGBColor(220, 220, 220)

    def __init__(self, output_dir: str = "./reports"):
        """
        初始化报告生成器。

        Parameters
        ----------
        output_dir : str
            报告输出目录
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """添加标题（1=一级，2=二级，3=三级）"""
        heading = doc.add_heading(text, level=level)

        # 统一设置中文字体
        for run in heading.runs:
            run.font.name = self.FONT_HEITI
            # 设置 East Asian 字体
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), self.FONT_HEITI)

        if level == 1:
            heading.style = "Heading 1"
            for run in heading.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = self.COLOR_BLACK
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif level == 2:
            heading.style = "Heading 2"
            for run in heading.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = self.COLOR_BLACK
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif level == 3:
            heading.style = "Heading 3"
            for run in heading.runs:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = self.COLOR_BLACK
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_paragraph(self, doc: Document, text: str, font_size: int = 12, bold: bool = False, alignment = None):
        """添加段落"""
        p = doc.add_paragraph(text)

        # 设置字体
        for run in p.runs:
            run.font.name = self.FONT_SONGTI
            run.font.size = Pt(font_size)
            # 设置 East Asian 字体
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), self.FONT_SONGTI)
            if bold:
                run.font.bold = True
            run.font.color.rgb = self.COLOR_BLACK

        # 设置段落格式
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        if alignment:
            p.alignment = alignment

        return p

    def _add_table(self, doc: Document, headers: List[str], rows: List[List[str]], col_widths: List[float] = None):
        """添加表格（3 列：项目、数值、来源）"""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = "Light Grid Accent 1"

        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                table.columns[i].width = Cm(width)

        # 设置表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header

            # 表头样式
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.FONT_SONGTI
                    run.font.size = Pt(10)
                    # 设置 East Asian 字体
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = OxmlElement('w:rFonts')
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), self.FONT_SONGTI)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 表头背景色（灰色）
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), "808080")
            cell._element.get_or_add_tcPr().append(shading_elm)

        # 填充数据行
        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_text

                # 单元格文本样式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = self.FONT_SONGTI
                        run.font.size = Pt(10)
                        # 设置 East Asian 字体
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.append(rFonts)
                        rFonts.set(qn('w:eastAsia'), self.FONT_SONGTI)
                        run.font.color.rgb = self.COLOR_BLACK
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_figure(self, doc: Document, fig, width_cm: float = 14.0):
        """嵌入一张子系统图件 + 居中图注（路径不存在则跳过）。"""
        import os
        path = getattr(fig, "path", None) or (fig.get("path") if isinstance(fig, dict) else None)
        caption = getattr(fig, "caption", None) or (fig.get("caption", "") if isinstance(fig, dict) else "")
        source = getattr(fig, "source", None) or (fig.get("source", "") if isinstance(fig, dict) else "")
        if not path or not os.path.exists(path):
            return
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Cm(width_cm))
        except Exception:
            return
        cap_text = f"图：{caption}" if caption else ""
        if not cap_text:
            return
        cap = self._add_paragraph(doc, cap_text, font_size=10)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.italic = True
            r.font.color.rgb = self.COLOR_GRAY

    # 置信等级 → 颜色
    _GRADE_COLORS = {
        "A": RGBColor(0, 150, 60),    # 绿
        "B": RGBColor(0, 110, 200),   # 蓝
        "C": RGBColor(220, 150, 0),   # 橙
        "D": RGBColor(200, 30, 30),   # 红
    }

    def _add_confidence_section(self, doc: Document, confidence: dict, mineral_type: str = ""):
        """渲染 A-B-C-D 综合置信评价（缺失时给出占位说明）。"""
        if not confidence:
            self._add_paragraph(doc, "[置信评价未生成] 综合研判环节未返回有效结果，建议复核各章节数据后重试。",
                                font_size=11)
            return

        grade = str(confidence.get("grade", "")).strip().upper()[:1]
        label = confidence.get("grade_label", "")
        # 大号醒目等级
        gp = doc.add_paragraph()
        gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = gp.add_run(f"综合置信等级：{grade or '—'}")
        run.font.name = self.FONT_HEITI
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = self._GRADE_COLORS.get(grade, self.COLOR_BLACK)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), self.FONT_HEITI)
        if label:
            lp = self._add_paragraph(doc, f"（{label}）", font_size=12)
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_paragraph(doc, "评价标准：A（最高/最有利）> B > C > D（最低）。", font_size=10)

        summary = confidence.get("summary", "")
        if summary:
            doc.add_paragraph()
            self._add_paragraph(doc, summary, font_size=11)

        dims = confidence.get("dimensions", [])
        if dims:
            doc.add_paragraph()
            self._add_heading(doc, "分项研判", level=3)
            rows = [[d.get("name", ""), d.get("level", ""), d.get("note", "")] for d in dims]
            self._add_table(doc, ["评价维度", "有利程度", "依据"], rows, col_widths=[4, 2.5, 7])

        rec = confidence.get("recommendation", "")
        if rec:
            doc.add_paragraph()
            self._add_heading(doc, "下一步工作建议", level=3)
            self._add_paragraph(doc, rec, font_size=11)

    def _add_front_matter(self, doc: Document, location: LocationContext,
                          search_results: Dict[str, SearchResult], mineral_type: str,
                          target_figure=None, confidence: dict = None):
        """前言（工作目标/数据构成/技术路线）+ 执行摘要（结论先行），实现报告首尾呼应。"""
        # ---- 前言 ----
        self._add_heading(doc, "前言", level=1)

        goal = (f"本报告针对研究区「{location.area_name}」开展面向"
                f"「{mineral_type}」的找矿前景综合评价；" if mineral_type
                else f"本报告针对研究区「{location.area_name}」开展区域地学综合评价；")
        goal += "通过多源资料的采集、结构化提取与综合研判，回答该区是否具备找矿/勘探价值、" \
                "有利方向何在、下一步如何部署等问题。"
        self._add_heading(doc, "一、工作目标", level=2)
        self._add_paragraph(doc, goal)

        # 数据构成：按来源层级统计章节数
        n_sub = n_api = n_web = 0
        for r in (search_results or {}).values():
            lv = getattr(r, "evidence_level", "") or ""
            if "子系统本地实证" in lv:
                n_sub += 1
            if "直连API" in lv:
                n_api += 1
            if "网络检索" in lv:
                n_web += 1
        self._add_heading(doc, "二、数据构成", level=2)
        self._add_paragraph(
            doc,
            f"本报告证据来自三个层级：子系统本地实证（{n_sub} 章涉及）、权威直连 API"
            f"（{n_api} 章涉及）、网络公开检索（{n_web} 章涉及）。各章末标注其来源层级，"
            "其中子系统本地实证可信度最高，仅网络检索者需补充查证。")

        self._add_heading(doc, "三、技术路线", level=2)
        self._add_paragraph(
            doc,
            "数据采集（直连 API + 缓存 + 网络检索 + 子系统实证）→ 结构化提取（各章定量证据）"
            "→ 综合研判（同一证据上下文统一产出靶区/有利地段评级与综合置信等级，二者自洽）"
            "→ 结论与下一步建议。")

        # ---- 执行摘要（结论先行）----
        doc.add_paragraph()
        self._add_heading(doc, "执行摘要（结论先行）", level=1)
        if confidence:
            grade = str(confidence.get("grade", "")).strip().upper()[:1]
            label = confidence.get("grade_label", "")
            gp = doc.add_paragraph()
            gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = gp.add_run(f"综合置信等级：{grade or '—'}" + (f"（{label}）" if label else ""))
            run.font.name = self.FONT_HEITI
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = self._GRADE_COLORS.get(grade, self.COLOR_BLACK)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), self.FONT_HEITI)

            summary = confidence.get("summary", "")
            if summary:
                self._add_paragraph(doc, summary, font_size=11)
            rec = confidence.get("recommendation", "")
            if rec:
                self._add_paragraph(doc, f"下一步建议：{rec}", font_size=11, bold=True)
        else:
            self._add_paragraph(doc, "（综合置信评价见报告末章。）", font_size=11)

        # 最有利方向（靶区/有利地段）
        if target_figure is not None:
            mode = getattr(target_figure, "mode", "targets")
            if mode == "areas":
                areas = getattr(target_figure, "favorable_areas", None) or []
                if areas:
                    top = areas[0]
                    self._add_paragraph(
                        doc, f"最有利方向：{top.get('name','')}（{top.get('direction','')}），"
                             f"详见「找矿有利地段」章。", font_size=11)
            else:
                targets = getattr(target_figure, "targets", None) or []
                if targets:
                    t = targets[0]
                    self._add_paragraph(
                        doc, f"首选靶区：#{t.get('rank','')}（{t.get('grade','')} 级，中心 "
                             f"{t.get('longitude',0):.4f}, {t.get('latitude',0):.4f}），"
                             f"详见「靶区推荐」章。", font_size=11)

        doc.add_page_break()

    def _add_key_findings(self, doc: Document, findings: List[str]):
        """添加关键发现列表（过滤无数据占位项）"""
        findings = [f for f in (findings or []) if f and not _is_no_data(f)]
        if not findings:
            return

        self._add_paragraph(doc, "主要发现：", font_size=12, bold=True)

        for finding in findings:
            # 使用符号列表
            p = doc.add_paragraph(finding, style="List Bullet")
            for run in p.runs:
                run.font.name = self.FONT_SONGTI
                run.font.size = Pt(11)
                # 设置 East Asian 字体
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:eastAsia'), self.FONT_SONGTI)
                run.font.color.rgb = self.COLOR_BLACK
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    def build_report(self, location: LocationContext, search_results: Dict[str, SearchResult],
                     output_name: str = None, mineral_type: str = "",
                     target_figure=None, confidence: dict = None) -> str:
        """
        生成完整的地质勘探报告。

        Parameters
        ----------
        location : LocationContext
            地理位置上下文
        search_results : Dict[str, SearchResult]
            所有搜索结果（类别 ID → SearchResult）
        output_name : str
            输出文件名（不含扩展名）。如果为 None，使用 area_name 和日期生成。

        Returns
        -------
        str
            生成的 .docx 文件路径
        """
        doc = Document()

        # 设置页面边距（GB/T 9704）
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        # === 封面页 ===
        self._add_paragraph(doc, "")  # 空行
        self._add_heading(doc, "地质勘探综合报告", level=1)
        self._add_paragraph(doc, "")

        # 区块信息
        info_lines = [
            f"研究区：{location.area_name}",
            "",
            f"位置：{location.location_str}（{location.country}）",
            f"坐标范围：{location.coords_str}",
        ]
        if mineral_type:
            info_lines.append(f"目标矿种：{mineral_type}")
        info_lines += ["", f"报告生成日期：{datetime.now().strftime('%Y年%m月%d日')}"]
        info_text = "\n".join(info_lines)

        self._add_paragraph(doc, info_text, font_size=12)

        # 分页
        doc.add_page_break()

        # === 前言 + 执行摘要（结论先行，首尾呼应）===
        self._add_front_matter(doc, location, search_results, mineral_type,
                               target_figure=target_figure, confidence=confidence)

        # === 第一章：研究区基本信息 ===
        self._add_heading(doc, "第一章  研究区基本信息", level=1)

        self._add_heading(doc, "1.1 地理位置", level=2)
        self._add_paragraph(doc, f"研究区位于 {location.location_str}，坐标范围为：{location.coords_str}。")

        self._add_heading(doc, "1.2 行政区划", level=2)
        admin_text = f"国家：{location.country}；省份：{location.province or '未确定'}；城市：{location.city or '未确定'}；区县：{location.district or '未确定'}。"
        self._add_paragraph(doc, admin_text)

        self._add_heading(doc, "1.3 地质背景", level=2)
        if location.kml_description:
            self._add_paragraph(doc, location.kml_description)
        else:
            self._add_paragraph(doc, "区块地质背景详见后续各章节描述。")

        # 分页
        doc.add_page_break()

        # === 第二章至第九章：8 类地学数据 ===
        for chapter_num, category in enumerate(get_all_categories(), start=2):
            cat_id = category.id
            result = search_results.get(cat_id)

            if not result:
                continue

            self._add_heading(doc, f"第{self._num_to_chinese(chapter_num)}章  {category.chapter_title}", level=1)

            if result.error:
                # 如果搜索失败，显示错误消息
                error_p = self._add_paragraph(doc, f"[数据获取失败] {result.error}", font_size=11)
                for run in error_p.runs:
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                # 添加概述（无公开数据的占位概述不展示）
                if result.summary and not _is_no_data(result.summary):
                    self._add_paragraph(doc, result.summary, font_size=11)
                    doc.add_paragraph()  # 空行

                # 添加数据表格（去掉来源列；过滤无公开数据的占位行）
                valid_dps = [dp for dp in result.data_points if not _is_no_data(dp.value)]
                if valid_dps:
                    table_rows = [[dp.item, dp.value] for dp in valid_dps]
                    self._add_table(doc, ["项目", "数值/描述"], table_rows, col_widths=[5, 9])
                    doc.add_paragraph()  # 空行

                # 添加关键发现
                if result.key_findings:
                    self._add_key_findings(doc, result.key_findings)

                # 添加勘探影响分析
                if result.exploration_impact and mineral_type:
                    doc.add_paragraph()
                    self._add_heading(doc, f"对{mineral_type}勘探作业的影响", level=3)
                    self._add_paragraph(doc, result.exploration_impact, font_size=11)

                # 嵌入子系统图件（蚀变/物探/深部预测等）
                _figs = getattr(result, "figures", None) or []
                if _figs:
                    doc.add_paragraph()
                    for fig in _figs:
                        self._add_figure(doc, fig)

                # 证据来源层级标注（供可信度核验；不展示具体网址/出处）
                level = getattr(result, "evidence_level", "")
                if level:
                    lp = self._add_paragraph(doc, f"数据来源层级：{level}", font_size=9)
                    for run in lp.runs:
                        run.font.italic = True
                        run.font.color.rgb = self.COLOR_GRAY

            # 每个章节后分页
            doc.add_page_break()

        # === 靶区推荐章 ===
        next_chapter = 2 + len(get_all_categories())
        mode = getattr(target_figure, "mode", "targets") if target_figure is not None else None
        if mode == "areas":
            self._add_heading(doc, f"第{self._num_to_chinese(next_chapter)}章  找矿有利地段", level=1)
        else:
            self._add_heading(doc, f"第{self._num_to_chinese(next_chapter)}章  靶区推荐", level=1)
        if target_figure is not None:
            if mode == "areas":
                self._add_paragraph(
                    doc,
                    "本研究区暂无深部探测靶区数据，下述结论为综合多源证据（地质、地球物理、"
                    "地球化学、遥感蚀变等）的定性研判，圈定若干找矿有利地段/远景区。"
                    "上述地段未经深部探测验证，具体靶区坐标待野外查证后厘定。",
                    font_size=11)
                self._add_figure(doc, target_figure, width_cm=15.0)
                areas = getattr(target_figure, "favorable_areas", None) or []
                if areas:
                    rows = []
                    for a in areas:
                        rows.append([
                            a.get("name", ""), a.get("direction", ""),
                            a.get("basis", ""), a.get("method", "")])
                    self._add_table(doc, ["有利地段", "方位/特征", "判断依据", "建议查证手段"], rows,
                                    col_widths=[2.6, 3.0, 5.4, 3.0])
            else:
                tgt_source = getattr(target_figure, "target_source", "exploration")
                if tgt_source == "drill":
                    intro = ("综合各方面资料（地质、地球物理、地球化学、遥感蚀变、深部探测及三维成矿建模），"
                             "采用 geo-drill 的 AI 钻探布孔作为推荐靶区（孔位即靶区，含真实坐标与目标深度），"
                             "在下图中以高热力弧形圈定，并对各靶区给出 A-B-C-D 置信评级（A 最高）。")
                else:
                    intro = ("综合各方面资料（地质、地球物理、地球化学、遥感蚀变及深部探测），在下图中以"
                             "高热力弧形圈定推荐找矿靶区，并对各靶区给出 A-B-C-D 置信评级（A 最高）。")
                self._add_paragraph(doc, intro, font_size=11)
                self._add_figure(doc, target_figure, width_cm=15.0)
                targets = getattr(target_figure, "targets", None)
                if targets:
                    rows = []
                    for t in targets:
                        coord = f"{t.get('longitude', 0):.4f}, {t.get('latitude', 0):.4f}"
                        rows.append([f"#{t.get('rank', '')}", t.get("grade", ""), coord, t.get("reason", "")])
                    self._add_table(doc, ["靶区", "置信等级", "中心坐标(°E, °N)", "评分理由"], rows,
                                    col_widths=[1.6, 2.0, 4.0, 8.0])

        # 已知矿点分布（deposits 本地实证）+ 钻探布孔/反馈（geo-drill 验证闭环），有数据才出现
        try:
            from .synthesis import get_known_deposits, get_drill_evidence
            deposits = get_known_deposits(location)
            drill = get_drill_evidence(location)
        except Exception as exc:  # noqa: BLE001
            print(f"[Report] 靶区验证数据获取失败：{exc}")
            deposits, drill = [], {}
        if deposits:
            self._add_paragraph(
                doc, "区内及周边已知矿点（真实标签，用于靶区空间套合与成矿模型印证）：", font_size=11)
            rows = [[d.get("name", "") or "(未命名)", d.get("commodity", ""), d.get("deposit_type", ""),
                     f"{d.get('lon', 0):.4f}, {d.get('lat', 0):.4f}"] for d in deposits[:15]]
            self._add_table(doc, ["矿点名称", "矿种", "矿床类型", "坐标(°E, °N)"], rows,
                            col_widths=[4.2, 2.2, 3.6, 4.0])
        holes = (drill.get("holes") or []) if drill else []
        feedback = (drill.get("feedback") or []) if drill else []
        # 若计划孔已作为主靶区表展示（target_source=drill），此处不再重复列出，仅保留反馈
        main_is_drill = getattr(target_figure, "target_source", None) == "drill"
        show_holes = holes and not main_is_drill
        if show_holes or feedback:
            self._add_paragraph(
                doc, "钻探验证与布孔建议（geo-drill 决策支持，需工程实施验证）：", font_size=11)
            if feedback:
                ore = sum(1 for f in feedback if f.get("outcome") == "ore")
                self._add_paragraph(
                    doc, f"已有钻孔反馈 {len(feedback)} 个：见矿 {ore}，无矿 {len(feedback)-ore}。", font_size=10)
                rows = [[f.get("hole_id", "") or "-", f.get("outcome", ""), f.get("element", ""),
                         str(f.get("max_grade", ""))] for f in feedback[:12]]
                self._add_table(doc, ["孔号", "结果", "元素", "最高品位"], rows,
                                col_widths=[3.0, 2.4, 2.4, 3.0])
            if show_holes:
                rows = [[h.get("hole_id", "") or f"#{h.get('rank', '')}",
                         f"{h.get('lon', 0):.4f}, {h.get('lat', 0):.4f}",
                         str(h.get("target_depth_m", "")), str(h.get("priority", h.get("score", "")))]
                        for h in holes[:12]]
                self._add_table(doc, ["计划孔", "坐标(°E, °N)", "目标深度(m)", "优先级/评分"], rows,
                                col_widths=[2.8, 4.0, 2.6, 3.0])
        doc.add_page_break()

        # === 综合置信评价章（A-B-C-D）===
        conf_chapter = next_chapter + 1
        self._add_heading(doc, f"第{self._num_to_chinese(conf_chapter)}章  综合置信评价", level=1)
        self._add_confidence_section(doc, confidence, mineral_type)

        # === 生成文件 ===
        if not output_name:
            timestamp = datetime.now().strftime("%Y%m%d")
            output_name = f"{location.area_name}_{timestamp}"

        output_path = self.output_dir / f"{output_name}.docx"
        doc.save(str(output_path))

        # 同时生成 PPT 演示文稿
        from .pptx_builder import PptxBuilder
        pptx_builder = PptxBuilder(str(self.output_dir))
        pptx_path = pptx_builder.build_pptx(location, search_results, output_name, mineral_type,
                                            target_figure=target_figure, confidence=confidence)

        return str(output_path), pptx_path

    @staticmethod
    def _num_to_chinese(num: int) -> str:
        """将数字转换为中文（一、二、三...）"""
        nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九"]
        if num < 10:
            return nums[num]
        elif num < 100:
            tens = num // 10
            ones = num % 10
            result = nums[tens] + "十"
            if ones > 0:
                result += nums[ones]
            return result
        else:
            return str(num)
